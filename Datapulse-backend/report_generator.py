"""
ML Analysis Report Generator
Generates comprehensive Word and PDF reports of the entire ML workflow
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import logging
import traceback

logger = logging.getLogger(__name__)


class MLReportGenerator:
    """Generate comprehensive ML analysis reports"""

    def __init__(self):
        self.doc = None

    def create_report(self, session_data, output_format='docx'):
        try:
            self.doc = Document()
            self._setup_styles()

            self._add_title_page(session_data)
            self._add_dataset_info(session_data)
            self._add_cleaning_operations(session_data)
            self._add_outlier_treatment(session_data)
            self._add_model_training(session_data)
            self._add_performance_metrics(session_data)
            self._add_feature_importance(session_data)
            self._add_prediction_example(session_data)
            self._add_insights(session_data)
            self._add_footer()

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            session_id = session_data.get('session_id', 'unknown')

            docx_path = f'/tmp/ml_analysis_report_{session_id}_{timestamp}.docx'
            self.doc.save(docx_path)
            logger.info(f"Report saved: {docx_path}")

            reports = {'docx': docx_path}

            if output_format in ['pdf', 'both']:
                try:
                    from pdf_generator import PDFReportGenerator
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    pdf_gen = PDFReportGenerator()
                    pdf_gen.create_pdf_report(session_data, pdf_path)
                    reports['pdf'] = pdf_path
                    logger.info(f"PDF report saved: {pdf_path}")
                except Exception as e:
                    logger.error(f"PDF conversion failed: {e}")
                    reports['pdf_error'] = str(e)

            return reports

        except Exception as e:
            logger.error(f"Report generation error: {e}")
            raise

    def _setup_styles(self):
        styles = self.doc.styles

        if 'Custom Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Calibri'
            heading1.font.size = Pt(24)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(68, 114, 196)

        if 'Custom Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.name = 'Calibri'
            heading2.font.size = Pt(18)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(68, 114, 196)

    def _add_title_page(self, data):
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('Machine Learning Analysis Report')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(68, 114, 196)

        self.doc.add_paragraph()

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Comprehensive Data Science Workflow Documentation')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(89, 89, 89)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
        info.add_run(f"Session ID: {data.get('session_id', 'N/A')}\n")
        if data.get('filename'):
            info.add_run(f"Dataset: {data.get('filename')}\n")

        self.doc.add_page_break()

    def _add_dataset_info(self, data):
        self.doc.add_heading('1. Dataset Overview', level=1)

        summary = data.get('summary', {})

        p = self.doc.add_paragraph()
        p.add_run('Dataset Statistics:\n').bold = True

        total_rows = summary.get('total_rows', 'N/A')
        if isinstance(total_rows, (int, float)):
            p.add_run(f"• Total Rows: {int(total_rows):,}\n")
        else:
            p.add_run(f"• Total Rows: {total_rows}\n")

        total_cols = summary.get('total_columns', 'N/A')
        if isinstance(total_cols, (int, float)):
            p.add_run(f"• Total Columns: {int(total_cols)}\n")
        else:
            p.add_run(f"• Total Columns: {total_cols}\n")

        p.add_run(f"• Numeric Columns: {summary.get('numeric_columns', 'N/A')}\n")
        p.add_run(f"• Categorical Columns: {summary.get('categorical_columns', 'N/A')}\n")
        p.add_run(f"• Missing Values: {summary.get('missing_values_total', 'N/A')}\n")
        p.add_run(f"• Duplicate Rows: {summary.get('duplicate_rows', 'N/A')}\n")

        # Column information table
        if summary.get('column_info'):
            self.doc.add_heading('Column Information', level=2)

            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            headers = ['Column', 'Type', 'Non-Null', 'Unique', 'Missing']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for col in summary['column_info'][:20]:
                row = table.add_row()
                row.cells[0].text = str(col['name'])[:30]
                row.cells[1].text = str(col['dtype'])
                row.cells[2].text = str(col['non_null'])
                row.cells[3].text = str(col['unique'])
                row.cells[4].text = str(col['missing'])

        self.doc.add_page_break()

    def _add_cleaning_operations(self, data):
        self.doc.add_heading('2. Data Cleaning Operations', level=1)

        cleaning = data.get('cleaning_operations', [])

        if not cleaning:
            self.doc.add_paragraph('No data cleaning operations were performed.')
        else:
            for i, op in enumerate(cleaning, 1):
                p = self.doc.add_paragraph(style='List Number')

                op_type = op.get('type', 'Unknown')
                if op_type == 'remove_duplicates':
                    p.add_run(f"Removed Duplicates: ").bold = True
                    p.add_run(f"{op.get('rows_removed', 0)} duplicate rows were removed")
                elif op_type == 'handle_missing':
                    p.add_run(f"Handled Missing Values in '{op.get('column')}': ").bold = True
                    method = op.get('method', 'N/A')
                    method_labels = {
                        'drop': 'Dropped rows with missing values',
                        'fill_mean': 'Filled with mean value',
                        'fill_median': 'Filled with median value',
                        'fill_mode': 'Filled with most frequent value',
                        'fill_forward': 'Forward filled values',
                        'fill_backward': 'Backward filled values',
                    }
                    p.add_run(method_labels.get(method, method))
                elif op_type == 'drop_column':
                    p.add_run(f"Dropped Column: ").bold = True
                    p.add_run(f"Removed column '{op.get('column')}'")
                elif op_type == 'rename_column':
                    p.add_run(f"Renamed Column: ").bold = True
                    p.add_run(f"'{op.get('old_name')}' → '{op.get('new_name')}'")
                else:
                    p.add_run(f"{op_type}")

        self.doc.add_page_break()

    def _add_outlier_treatment(self, data):
        self.doc.add_heading('3. Outlier Treatment', level=1)

        outliers = data.get('outlier_treatment', [])

        if not outliers:
            self.doc.add_paragraph('No outlier treatment was performed.')
        else:
            for op in outliers:
                p = self.doc.add_paragraph(style='List Bullet')
                column = op.get('column', 'Unknown')
                method = op.get('method', 'Unknown')
                outliers_found = op.get('outliers_found', 0)

                p.add_run(f"Column '{column}': ").bold = True
                if method == 'remove':
                    p.add_run(f"Removed {outliers_found} outliers")
                elif method == 'cap':
                    p.add_run(f"Capped {outliers_found} outliers at IQR boundaries")
                elif method == 'transform':
                    p.add_run(f"Applied {op.get('transform_type', 'log')} transformation")
                else:
                    p.add_run(f"Method: {method}")

        self.doc.add_page_break()

    def _add_model_training(self, data):
        self.doc.add_heading('4. Model Training Configuration', level=1)

        model_config = data.get('model_config', {})

        if not model_config or not model_config.get('task_type'):
            self.doc.add_paragraph('No model training was performed.')
            return

        p = self.doc.add_paragraph()
        p.add_run('Training Configuration:\n').bold = True

        task_type = model_config.get('task_type', 'N/A')
        p.add_run(f"• Task Type: {task_type.title() if isinstance(task_type, str) else task_type}\n")
        p.add_run(f"• Target Column: {model_config.get('target_column', 'N/A')}\n")
        p.add_run(f"• Algorithm: {model_config.get('model_type', 'N/A')}\n")

        test_size = model_config.get('test_size', 0.2)
        if isinstance(test_size, (int, float)):
            p.add_run(f"• Test Size: {test_size * 100:.0f}%\n")
            p.add_run(f"• Train Size: {(1 - test_size) * 100:.0f}%\n")
        else:
            p.add_run(f"• Test Size: {test_size}\n")

        p.add_run(f"• Hyperparameter Tuning: {'Enabled (GridSearchCV)' if model_config.get('tune_params') else 'Disabled (Default Parameters)'}\n")

        if model_config.get('trained_at'):
            p.add_run(f"• Training Completed: {model_config.get('trained_at')}\n")

    def _add_performance_metrics(self, data):
        self.doc.add_heading('5. Model Performance Metrics', level=1)

        report = data.get('performance_report', {})
        task_type = data.get('model_config', {}).get('task_type', 'classification')

        if not report:
            self.doc.add_paragraph('No performance metrics available.')
            return

        if task_type == 'classification':
            self._add_classification_metrics(report)
        else:
            self._add_regression_metrics(report)

    def _add_classification_metrics(self, report):
        p = self.doc.add_paragraph()

        accuracy = report.get('accuracy', 0)
        p.add_run(f"Accuracy: ").bold = True
        p.add_run(f"{accuracy * 100:.2f}%\n")

        f1 = report.get('F1_Score', 0)
        p.add_run(f"F1 Score: ").bold = True
        p.add_run(f"{f1:.4f}\n")

        if 'weighted avg' in report:
            weighted = report['weighted avg']
            p.add_run(f"Precision (Weighted): ").bold = True
            p.add_run(f"{weighted.get('precision', 0):.4f}\n")
            p.add_run(f"Recall (Weighted): ").bold = True
            p.add_run(f"{weighted.get('recall', 0):.4f}\n")

        cv_score = report.get('Cross_Validation_Score', 0)
        p.add_run(f"Cross-Validation Score: ").bold = True
        p.add_run(f"{cv_score:.4f}\n")

        if 'ROC_AUC' in report:
            p.add_run(f"ROC AUC Score: ").bold = True
            p.add_run(f"{report['ROC_AUC']:.4f}\n")

        self.doc.add_paragraph()
        self.doc.add_heading('Per-Class Metrics', level=2)

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        headers = ['Class', 'Precision', 'Recall', 'F1-Score']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for key, value in report.items():
            if isinstance(value, dict) and 'precision' in value:
                if key not in ['accuracy', 'macro avg', 'weighted avg']:
                    row = table.add_row()
                    row.cells[0].text = str(key)
                    row.cells[1].text = f"{value.get('precision', 0):.4f}"
                    row.cells[2].text = f"{value.get('recall', 0):.4f}"
                    row.cells[3].text = f"{value.get('f1-score', 0):.4f}"

    def _add_regression_metrics(self, report):
        p = self.doc.add_paragraph()

        r2 = report.get('R² Score', 0)
        p.add_run(f"R² Score: ").bold = True
        p.add_run(f"{r2:.4f} ({r2 * 100:.1f}% variance explained)\n")

        mse = report.get('Mean Squared Error', 0)
        p.add_run(f"Mean Squared Error: ").bold = True
        p.add_run(f"{mse:.4f}\n")

        rmse = mse ** 0.5
        p.add_run(f"Root Mean Squared Error: ").bold = True
        p.add_run(f"{rmse:.4f}\n")

        mae = report.get('Mean Absolute Error', 0)
        p.add_run(f"Mean Absolute Error: ").bold = True
        p.add_run(f"{mae:.4f}\n")

        cv_score = report.get('Cross_Validation_Score', 0)
        p.add_run(f"Cross-Validation Score: ").bold = True
        p.add_run(f"{cv_score:.4f}\n")

    def _add_feature_importance(self, data):
        self.doc.add_page_break()
        self.doc.add_heading('6. Feature Importance Analysis', level=1)

        feature_importance = data.get('performance_report', {}).get('Feature_Importance', [])

        if not feature_importance:
            self.doc.add_paragraph('Feature importance data not available for this model.')
        else:
            self.doc.add_paragraph('Top 10 Most Important Features:')

            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            headers = ['Rank', 'Feature', 'Importance Score']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for i, feature in enumerate(feature_importance[:10], 1):
                row = table.add_row()
                row.cells[0].text = f"#{i}"
                row.cells[1].text = feature.get('Feature', 'Unknown')
                row.cells[2].text = f"{feature.get('Importance', 0):.6f}"

    def _add_prediction_example(self, data):
        self.doc.add_page_break()
        self.doc.add_heading('7. Example Prediction', level=1)

        prediction = data.get('example_prediction')

        if not prediction:
            self.doc.add_paragraph('No example prediction was performed.')
        else:
            self.doc.add_heading('Input Values:', level=2)

            inputs = prediction.get('inputs', {})
            for feature, value in inputs.items():
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{feature}: ").bold = True
                p.add_run(str(value))

            self.doc.add_paragraph()
            self.doc.add_heading('Prediction Result:', level=2)

            result = prediction.get('result', {})
            pred_value = result.get('prediction', ['N/A'])[0]

            p = self.doc.add_paragraph()
            p.add_run('Predicted Value: ').bold = True
            run = p.add_run(str(pred_value))
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 128, 0)

            if 'probability' in result:
                self.doc.add_paragraph()
                self.doc.add_heading('Class Probabilities:', level=2)
                probs = result['probability']
                for class_name, prob in probs.items():
                    p = self.doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{class_name}: ").bold = True
                    p.add_run(f"{prob * 100:.2f}%")

    def _add_insights(self, data):
        self.doc.add_page_break()
        self.doc.add_heading('8. AI-Generated Insights', level=1)

        insights = data.get('insights', {})

        if not insights:
            self.doc.add_paragraph('No insights were generated for this analysis.')
        else:
            if insights.get('key_findings'):
                self.doc.add_heading('Key Findings:', level=2)
                for finding in insights['key_findings']:
                    self.doc.add_paragraph(finding, style='List Bullet')

            if insights.get('recommendations'):
                self.doc.add_paragraph()
                self.doc.add_heading('Recommendations:', level=2)
                for rec in insights['recommendations']:
                    self.doc.add_paragraph(rec, style='List Bullet')

            if insights.get('data_quality'):
                self.doc.add_paragraph()
                self.doc.add_heading('Data Quality Notes:', level=2)
                for note in insights['data_quality']:
                    self.doc.add_paragraph(note, style='List Bullet')

    def _add_footer(self):
        self.doc.add_page_break()

        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('End of Report')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        footer2 = self.doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = footer2.add_run(f'Generated by DataPulse ML Platform • {datetime.now().year}')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(128, 128, 128)


# ============================================================================
# Wrapper functions for app.py compatibility
# ============================================================================

def build_docx(report_data: dict) -> bytes:
    """
    Build a DOCX report from the report_data dictionary.
    Returns the document as bytes.
    """
    try:
        session_data = {
            'session_id': report_data.get('upload', {}).get('session_id', 'unknown'),
            'filename': report_data.get('upload', {}).get('filename', 'dataset.csv'),
            'summary': report_data.get('summary', {}),
            'cleaning_operations': _transform_cleaning(report_data.get('cleaning', {})),
            'outlier_treatment': report_data.get('outliers', []),
            'model_config': report_data.get('ml', {}),
            'performance_report': report_data.get('ml', {}).get('report', {}),
            'example_prediction': report_data.get('prediction'),
            'insights': report_data.get('insights', {})
        }

        generator = MLReportGenerator()
        generator.doc = Document()
        generator._setup_styles()

        generator._add_title_page(session_data)
        generator._add_dataset_info(session_data)
        generator._add_cleaning_operations(session_data)
        generator._add_outlier_treatment(session_data)
        generator._add_model_training(session_data)
        generator._add_performance_metrics(session_data)
        generator._add_feature_importance(session_data)
        generator._add_prediction_example(session_data)
        generator._add_insights(session_data)
        generator._add_footer()

        docx_buffer = io.BytesIO()
        generator.doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer.getvalue()

    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        logger.error(traceback.format_exc())
        raise


def build_pdf(report_data: dict) -> bytes:
    """
    Build a PDF report from the report_data dictionary.
    Returns the PDF as bytes.
    """
    try:
        from pdf_generator import PDFReportGenerator

        session_data = {
            'session_id': report_data.get('upload', {}).get('session_id', 'unknown'),
            'filename': report_data.get('upload', {}).get('filename', 'dataset.csv'),
            'summary': report_data.get('summary', {}),
            'cleaning_operations': _transform_cleaning(report_data.get('cleaning', {})),
            'outlier_treatment': report_data.get('outliers', []),
            'model_config': report_data.get('ml', {}),
            'performance_report': report_data.get('ml', {}).get('report', {}),
            'example_prediction': report_data.get('prediction'),
            'insights': report_data.get('insights', {})
        }

        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
            pdf_path = pdf_file.name

        pdf_gen = PDFReportGenerator()
        pdf_gen.create_pdf_report(session_data, pdf_path)

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()

        os.unlink(pdf_path)
        return pdf_bytes

    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        logger.error(traceback.format_exc())
        raise


# ============================================================================
# Helper functions
# ============================================================================

def _transform_cleaning(cleaning: dict) -> list:
    """Transform cleaning data to operations list"""
    operations = []

    if 'auto_clean' in cleaning:
        auto = cleaning['auto_clean']
        if auto.get('duplicates_removed', 0) > 0:
            operations.append({
                'type': 'remove_duplicates',
                'rows_removed': auto['duplicates_removed']
            })
        for col, method in auto.get('missing_handled', {}).items():
            operations.append({
                'type': 'handle_missing',
                'column': col,
                'method': method
            })

    if 'manual' in cleaning:
        manual = cleaning['manual']
        if manual.get('remove_duplicates'):
            operations.append({
                'type': 'remove_duplicates',
                'rows_removed': 0
            })
        for col, method in manual.get('missing_actions', {}).items():
            operations.append({
                'type': 'handle_missing',
                'column': col,
                'method': method
            })

    return operations
